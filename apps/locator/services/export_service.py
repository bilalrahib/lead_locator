import csv
import io
import pandas as pd
from docx import Document
from docx.shared import Inches
from typing import List
from django.http import HttpResponse
from ..models import SearchHistory, LocationData
import logging

logger = logging.getLogger(__name__)


class ExportService:
    """
    Service for exporting search results to various formats.
    """
    
    @staticmethod
    def export_to_csv(search_history: SearchHistory) -> HttpResponse:
        """
        Export search results to CSV format.
        
        Args:
            search_history: SearchHistory instance
            
        Returns:
            HttpResponse with CSV data
        """
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="vending_locations_{search_history.zip_code}.csv"'
        
        writer = csv.writer(response)
        
        # Write header
        writer.writerow([
            'Business Name',
            'Address',
            'Phone',
            'Email',
            'Website',
            'Google Rating',
            'Total Reviews',
            'Business Hours',
            'Foot Traffic Estimate',
            'Category',
            'Contact Completeness',
            'Priority Score',
            'Google Maps URL',
            'Latitude',
            'Longitude',
        ])
        
        # Write data rows
        for location in search_history.locations.all().order_by('-priority_score'):
            writer.writerow([
                location.name,
                location.address,
                location.phone,
                location.email,
                location.website,
                location.google_rating or '',
                location.google_user_ratings_total or '',
                location.business_hours_text.replace('\n', '; '),
                location.get_foot_traffic_estimate_display(),
                location.detailed_category,
                location.get_contact_completeness_display(),
                location.priority_score,
                location.maps_url,
                location.latitude,
                location.longitude,
            ])
        
        return response
    
    @staticmethod
    def export_to_xlsx(search_history: SearchHistory) -> HttpResponse:
        """
        Export search results to Excel format.
        
        Args:
            search_history: SearchHistory instance
            
        Returns:
            HttpResponse with Excel data
        """
        # Prepare data
        data = []
        for location in search_history.locations.all().order_by('-priority_score'):
            data.append({
                'Business Name': location.name,
                'Address': location.address,
                'Phone': location.phone,
                'Email': location.email,
                'Website': location.website,
                'Google Rating': location.google_rating or '',
                'Total Reviews': location.google_user_ratings_total or '',
                'Business Hours': location.business_hours_text.replace('\n', '; '),
                'Foot Traffic Estimate': location.get_foot_traffic_estimate_display(),
                'Category': location.detailed_category,
                'Contact Completeness': location.get_contact_completeness_display(),
                'Priority Score': location.priority_score,
                'Google Maps URL': location.maps_url,
                'Latitude': float(location.latitude),
                'Longitude': float(location.longitude),
            })
        
        # Create DataFrame
        df = pd.DataFrame(data)
        
        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Vending Locations', index=False)
            
            # Get the workbook and worksheet
            workbook = writer.book
            worksheet = writer.sheets['Vending Locations']
            
            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="vending_locations_{search_history.zip_code}.xlsx"'
        
        return response
    
    @staticmethod
    def export_to_docx(search_history: SearchHistory) -> HttpResponse:
        """
        Export search results to Word document format.
        
        Args:
            search_history: SearchHistory instance
            
        Returns:
            HttpResponse with Word document data
        """
        document = Document()
        
        # Add title
        title = document.add_heading('Vending Machine Location Report', 0)
        
        # Add search details
        document.add_heading('Search Details', level=1)
        details_para = document.add_paragraph()
        details_para.add_run('Search Location: ').bold = True
        details_para.add_run(f'{search_history.zip_code}\n')
        details_para.add_run('Machine Type: ').bold = True
        details_para.add_run(f'{search_history.get_machine_type_display()}\n')
        details_para.add_run('Search Radius: ').bold = True
        details_para.add_run(f'{search_history.radius} miles\n')
        details_para.add_run('Total Results: ').bold = True
        details_para.add_run(f'{search_history.results_count}\n')
        details_para.add_run('Search Date: ').bold = True
        details_para.add_run(f'{search_history.created_at.strftime("%B %d, %Y")}\n')
        
        # Add results
        document.add_heading('Location Results', level=1)
        
        for i, location in enumerate(search_history.locations.all().order_by('-priority_score'), 1):
            # Location heading
            document.add_heading(f'{i}. {location.name}', level=2)
            
            # Location details
            details_para = document.add_paragraph()
            details_para.add_run('Address: ').bold = True
            details_para.add_run(f'{location.address}\n')
            
            if location.phone:
                details_para.add_run('Phone: ').bold = True
                details_para.add_run(f'{location.phone}\n')
            
            if location.email:
                details_para.add_run('Email: ').bold = True
                details_para.add_run(f'{location.email}\n')
            
            if location.website:
                details_para.add_run('Website: ').bold = True
                details_para.add_run(f'{location.website}\n')
            
            if location.google_rating:
                details_para.add_run('Google Rating: ').bold = True
                details_para.add_run(f'{location.google_rating}/5.0 ({location.google_user_ratings_total or 0} reviews)\n')
            
            details_para.add_run('Foot Traffic Estimate: ').bold = True
            details_para.add_run(f'{location.get_foot_traffic_estimate_display()}\n')
            
            details_para.add_run('Priority Score: ').bold = True
            details_para.add_run(f'{location.priority_score}\n')
            
            if location.business_hours_text:
                details_para.add_run('Business Hours: ').bold = True
                details_para.add_run(f'\n{location.business_hours_text}\n')
            
            if location.maps_url:
                details_para.add_run('Google Maps: ').bold = True
                details_para.add_run(f'{location.maps_url}\n')
            
            # Add separator
            document.add_paragraph('_' * 50)
        
        # Save to memory
        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)
        
        response = HttpResponse(
            doc_io.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="vending_locations_{search_history.zip_code}.docx"'
        
        return response